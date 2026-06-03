"""
将一篇结构化的中文个人陈述导出为符合保研投递规范的 .docx：
- 正文：中文宋体 + 西文 Times New Roman + 小四（14pt）
- 段落：首行缩进 2 字符、1.25 倍行距
- 小标题：加粗、与正文同号、段前/段后留白
- 页边距：上下左右各 1 英寸
- 自动写入文档属性（标题/作者/学科）

使用方式：把 SECTIONS 替换为自己的内容，运行 `python build_statement.py`，
即可在当前目录生成 OUTPUT 指定的 docx 文件。
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


TITLE = "个人陈述"
OUTPUT = "目标学校_目标学院_个人陈述_申请人姓名.docx"
AUTHOR = "申请人姓名"
SUBJECT = "目标学校 目标学院 申请个人陈述"

SECTIONS = [
    (
        "一、专业学习情况",
        "在此填写本科院校、专业、核心课程、GPA / 排名 / 平均分、英语成绩、"
        "以及自己所掌握的研究方法与工具（如 Stata、Python、R、LaTeX 等），"
        "强调学科训练如何塑造你的问题意识与分析路径。",
    ),
    (
        "二、学术背景与研究训练",
        "在此呈现科研经历：参与的项目、承担的角色、使用的数据库与方法、"
        "已完成的论文与会议成果。优先用事实和数据说话，避免堆砌形容词。",
    ),
    (
        "三、研究兴趣",
        "在此说明研究兴趣聚焦的方向及其与目标专业/导师的衔接，"
        "解释为什么这些问题值得研究、你为什么具备相应的研究基础。",
    ),
    (
        "四、硕博阶段学习与研究计划",
        "在此说明在目标学校/学院希望完成的学习目标、希望深化的方法训练、"
        "希望聚焦的研究议题，以及可能的研究路径与预期产出。",
    ),
    (
        "五、实践基础与就业目标",
        "在此补充课外实践、实习、调研与社会服务经历，"
        "并说明长期就业方向与学术训练之间的对应关系。",
    ),
]


def set_font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Inches(0.28)


def set_heading(paragraph):
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run(TITLE)
    set_font(title_run, 14, True)

    for heading, body in SECTIONS:
        h = doc.add_paragraph()
        set_heading(h)
        r = h.add_run(heading)
        set_font(r, 14, True)

        p = doc.add_paragraph()
        set_paragraph_format(p)
        br = p.add_run(body)
        set_font(br, 14)

    core = doc.core_properties
    core.title = TITLE
    core.author = AUTHOR
    core.subject = SUBJECT

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
