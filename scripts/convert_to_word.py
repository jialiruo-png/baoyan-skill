#!/usr/bin/env python3
"""Convert a Chinese personal statement in Markdown/plain text to a formatted .docx.

Default formatting rules:
- Main title: 宋体, bold, 小二 (18 pt)
- Section headings: 黑体, 小四 (12 pt)
- Title and section headings line spacing: 1.5
- Body text: 宋体, 五号 (10.5 pt)
- Body line spacing: 1.2

Supported input:
- Markdown or plain text
- Recognizes the main title from a line like '# 个人陈述' or a plain first line '个人陈述'
- Recognizes section headings from lines like '## 一、...' or plain lines matching '^一、'
- Preserves paragraphs separated by blank lines
- Strips simple markdown markers like leading '#' and surrounding '**'

Usage:
    python convert_to_word.py
    python convert_to_word.py input.md
    python convert_to_word.py input.md output.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception as exc:
    raise SystemExit(
        "Failed to import python-docx. Please install the `python-docx` package "
        "instead of the legacy `docx` package."
    ) from exc

TITLE_TEXT = "个人陈述"
TITLE_FONT = "宋体"
HEADING_FONT = "黑体"
BODY_FONT = "宋体"

TITLE_SIZE_PT = 18.0   # 小二
HEADING_SIZE_PT = 12.0 # 小四
BODY_SIZE_PT = 10.5    # 五号

TITLE_LINE_SPACING = 1.5
HEADING_LINE_SPACING = 1.5
BODY_LINE_SPACING = 1.2

SECTION_HEADING_RE = re.compile(r"^(?:##\s*)?([一二三四五六七八九十百]+、.+)$")
MAIN_TITLE_RE = re.compile(r"^(?:#\s*)?个人陈述\s*$")
MARKDOWN_TITLE_RE = re.compile(r"^#\s+(.+?)\s*$")
SALUTATION_RE = re.compile(r"^尊敬的.+老师：\s*$")
PREFERRED_NAME_RE = re.compile(r"(?:^|[_\-\s])(ps|statement)(?:$|[_\-\s])", re.IGNORECASE)

SKIP_PARTS = {".venv", ".git", "references", "scripts", "outputs", "__pycache__"}
SKIP_FILES = {"SKILL.md", "input-notes.md"}


def set_east_asia_font(run, font_name: str) -> None:
    """Set Western and East Asian fonts for a run."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def ensure_rfonts(style_or_normalized) -> OxmlElement:
    """Return an rFonts element, creating the XML nodes when needed."""
    rpr = style_or_normalized._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    return rfonts


def set_paragraph_line_spacing(paragraph, multiple: float) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = multiple
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def ensure_style(document: Document, style_name: str, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    styles = document.styles
    if style_name in styles:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rfonts = ensure_rfonts(style)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def clean_inline_markdown(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^#+\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = text.replace("`", "")
    return text.strip()


def split_blocks(raw: str) -> List[str]:
    normalized = raw.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return []
    blocks = re.split(r"\n\s*\n+", normalized)
    return [block.strip() for block in blocks if block.strip()]


def classify_block(block: str) -> str:
    one_line = block.replace("\n", " ").strip()
    if MAIN_TITLE_RE.match(one_line) or MARKDOWN_TITLE_RE.match(one_line):
        return "title"
    if SECTION_HEADING_RE.match(one_line):
        return "heading"
    if SALUTATION_RE.match(one_line):
        return "salutation"
    if one_line in {"此致", "敬礼", "敬礼！"}:
        return "signoff"
    return "body"


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph(style="PS Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(clean_inline_markdown(text) or TITLE_TEXT)
    set_east_asia_font(run, TITLE_FONT)
    run.font.size = Pt(TITLE_SIZE_PT)
    run.font.bold = True
    set_paragraph_line_spacing(p, TITLE_LINE_SPACING)


def add_heading(document: Document, text: str) -> None:
    cleaned = clean_inline_markdown(text)
    m = SECTION_HEADING_RE.match(cleaned)
    heading_text = m.group(1) if m else cleaned
    p = document.add_paragraph(style="PS Heading")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(heading_text)
    set_east_asia_font(run, HEADING_FONT)
    run.font.size = Pt(HEADING_SIZE_PT)
    run.font.bold = False
    set_paragraph_line_spacing(p, HEADING_LINE_SPACING)


def add_body_paragraph(document: Document, text: str, *, center: bool = False) -> None:
    p = document.add_paragraph(style="PS Body")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(clean_inline_markdown(text))
    set_east_asia_font(run, BODY_FONT)
    run.font.size = Pt(BODY_SIZE_PT)
    run.font.bold = False
    set_paragraph_line_spacing(p, BODY_LINE_SPACING)
    if not center:
        p.paragraph_format.first_line_indent = Pt(BODY_SIZE_PT * 2)
    else:
        p.paragraph_format.first_line_indent = Pt(0)


def add_multiline_body(document: Document, block: str) -> None:
    lines = [clean_inline_markdown(x) for x in block.splitlines() if clean_inline_markdown(x)]
    if not lines:
        return

    if len(lines) == 1:
        add_body_paragraph(document, lines[0])
        return

    if any(x in {"此致", "敬礼", "敬礼！"} for x in lines):
        for line in lines:
            add_body_paragraph(document, line)
        return

    add_body_paragraph(document, "".join(lines))


def format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    rfonts = ensure_rfonts(normal)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)

    ensure_style(document, "PS Title", TITLE_FONT, TITLE_SIZE_PT, bold=True)
    ensure_style(document, "PS Heading", HEADING_FONT, HEADING_SIZE_PT, bold=False)
    ensure_style(document, "PS Body", BODY_FONT, BODY_SIZE_PT, bold=False)


def convert_text_to_docx(input_path: Path, output_path: Path) -> None:
    raw = read_text_file(input_path)
    blocks = split_blocks(raw)

    document = Document()
    format_document(document)

    has_explicit_title = any(classify_block(block) == "title" for block in blocks[:2])
    if not has_explicit_title:
        add_title(document, TITLE_TEXT)

    for block in blocks:
        kind = classify_block(block)
        if kind == "title":
            add_title(document, block)
        elif kind == "heading":
            add_heading(document, block)
        else:
            add_multiline_body(document, block)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def is_candidate_text_file(path: Path) -> bool:
    if path.suffix.lower() not in {".md", ".txt"}:
        return False
    if any(part in SKIP_PARTS for part in path.parts):
        return False
    if path.name in SKIP_FILES:
        return False
    return True


def read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="utf-8-sig")


def score_candidate(path: Path) -> tuple[int, int]:
    score = 0
    lowered = path.stem.lower()
    text = read_text_file(path)
    first_lines = "\n".join(text.splitlines()[:5])

    if "ps" in lowered:
        score += 5
    if "陈述" in path.stem or "statement" in lowered:
        score += 4
    if PREFERRED_NAME_RE.search(lowered):
        score += 3
    if MAIN_TITLE_RE.search(first_lines):
        score += 5
    if "申请" in text and "专业" in text:
        score += 2
    if len(text.strip()) >= 800:
        score += 2
    if len(re.findall(r"[\u4e00-\u9fff]", text)) >= 500:
        score += 2

    return score, path.stat().st_mtime_ns


def find_latest_ps_text(base_dir: Path) -> Path:
    candidates = [path for path in base_dir.rglob("*") if path.is_file() and is_candidate_text_file(path)]
    if not candidates:
        raise SystemExit("No candidate PS text files found. Provide an input .md/.txt file explicitly.")

    ranked = [(score_candidate(path), path) for path in candidates]
    best_score = max(score for score, _ in ranked)
    if best_score[0] <= 0:
        raise SystemExit("Could not identify a likely PS text file automatically. Please pass the input file explicitly.")

    best_paths = [path for score, path in ranked if score == best_score]
    return best_paths[0]


def default_output_path(input_path: Path, base_dir: Path) -> Path:
    return base_dir / "outputs" / f"{input_path.stem}.docx"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert a personal statement in Markdown/plain text to formatted DOCX.")
    parser.add_argument("input", type=Path, nargs="?", help="Input .md or .txt file. Omit to auto-detect the latest PS text.")
    parser.add_argument("output", type=Path, nargs="?", help="Output .docx path. Defaults to outputs/<input_stem>.docx")
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    base_dir = Path.cwd()
    input_path = args.input if args.input is not None else find_latest_ps_text(base_dir)
    if not input_path.is_absolute():
        input_path = (base_dir / input_path).resolve()

    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}")

    output_path = args.output if args.output is not None else default_output_path(input_path, base_dir)
    if not output_path.is_absolute():
        output_path = (base_dir / output_path).resolve()

    if output_path.suffix.lower() != ".docx":
        raise SystemExit("Output file must end with .docx")

    convert_text_to_docx(input_path, output_path)
    print(f"Input: {input_path}")
    print(f"Created: {output_path}")
